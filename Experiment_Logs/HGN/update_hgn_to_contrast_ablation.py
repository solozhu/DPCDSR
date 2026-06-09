from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parent
WORKBOOK = Path(r"D:\Eruxim\GradProj\contrast_ablation.xlsx")
OUT_CSV = ROOT / "hgn_best_metrics.csv"

METRICS = ["MRR", "NDCG@5", "NDCG@10", "HR@1", "HR@5", "HR@10"]
DATASET_TO_ROW = {
    "Food-Kitchen": 74,
    "Movie-Book": 82,
    "Entertainment-Education": 90,
}
DOMAIN_TO_START_COL = {"x": 5, "y": 11}

METRIC_RE = re.compile(
    r"epoch=(?P<epoch>\d+)\s+(?P<split>valid|test)"
    r"\[MRR=(?P<MRR>[-+0-9.]+)\s+NDCG@5=(?P<NDCG5>[-+0-9.]+)\s+"
    r"NDCG@10=(?P<NDCG10>[-+0-9.]+)\s+HR@1=(?P<HR1>[-+0-9.]+)\s+"
    r"HR@5=(?P<HR5>[-+0-9.]+)\s+HR@10=(?P<HR10>[-+0-9.]+)\]"
)


def parse_filename(path: Path) -> tuple[str, str]:
    stem = path.stem
    dataset, domain = stem.rsplit(" ", 1)
    if dataset not in DATASET_TO_ROW or domain not in DOMAIN_TO_START_COL:
        raise ValueError(f"Unexpected HGN log filename: {path.name}")
    return dataset, domain


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def parse_metric_rows(path: Path) -> list[dict]:
    dataset, domain = parse_filename(path)
    rows = []
    for match in METRIC_RE.finditer(docx_text(path)):
        rows.append(
            {
                "project": "HGN",
                "dataset": dataset,
                "domain": domain,
                "split": match.group("split"),
                "epoch": int(match.group("epoch")),
                "source": str(path),
                "MRR": float(match.group("MRR")),
                "NDCG@5": float(match.group("NDCG5")),
                "NDCG@10": float(match.group("NDCG10")),
                "HR@1": float(match.group("HR1")),
                "HR@5": float(match.group("HR5")),
                "HR@10": float(match.group("HR10")),
            }
        )
    return rows


def choose_best(rows: list[dict]) -> dict:
    valid_rows = [row for row in rows if row["split"] == "valid"]
    test_by_epoch = {row["epoch"]: row for row in rows if row["split"] == "test"}
    if not valid_rows:
        raise RuntimeError("No validation rows found")
    for valid in sorted(valid_rows, key=lambda row: row["MRR"], reverse=True):
        test = test_by_epoch.get(valid["epoch"])
        if test:
            chosen = dict(test)
            chosen["selected_by_valid_epoch"] = valid["epoch"]
            chosen["valid_MRR"] = valid["MRR"]
            return chosen
    raise RuntimeError("No test row matching a validation-best epoch")


def collect_best() -> dict[tuple[str, str], dict]:
    best = {}
    for path in sorted(ROOT.glob("*.docx")):
        rows = parse_metric_rows(path)
        chosen = choose_best(rows)
        best[(chosen["dataset"], chosen["domain"])] = chosen
    missing = [
        (dataset, domain)
        for dataset in DATASET_TO_ROW
        for domain in DOMAIN_TO_START_COL
        if (dataset, domain) not in best
    ]
    if missing:
        raise RuntimeError(f"Missing HGN logs for {missing}")
    return best


def write_csv(best: dict[tuple[str, str], dict]) -> None:
    fields = [
        "project",
        "dataset",
        "domain",
        "split",
        "epoch",
        "selected_by_valid_epoch",
        "valid_MRR",
        *METRICS,
        "source",
    ]
    with OUT_CSV.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for key in sorted(best):
            row = best[key]
            writer.writerow({field: row.get(field, "") for field in fields})


def update_workbook(best: dict[tuple[str, str], dict]) -> None:
    wb = load_workbook(WORKBOOK)
    ws = wb["Sheet1"]
    for (dataset, domain), row in best.items():
        row_idx = DATASET_TO_ROW[dataset]
        start_col = DOMAIN_TO_START_COL[domain]
        for offset, metric in enumerate(METRICS):
            ws.cell(row_idx, start_col + offset).value = round(float(row[metric]), 4)
    wb.save(WORKBOOK)


def main() -> None:
    best = collect_best()
    write_csv(best)
    update_workbook(best)
    print(f"Updated {WORKBOOK}")
    print(f"Wrote {OUT_CSV}")
    for key in sorted(best):
        row = best[key]
        print(key, row["epoch"], [round(row[m], 4) for m in METRICS])


if __name__ == "__main__":
    main()
